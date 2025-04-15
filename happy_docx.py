from mcp.server.fastmcp import FastMCP
import sys
import os
import re
from docx import Document
import PyPDF2
from typing import Dict, List, Any, Union, Optional
import json


mcp = FastMCP("docx-filesystem")


allowed_paths = sys.argv[1:] if len(sys.argv) > 1 else ['.']

def is_path_allowed(path: str) -> bool:
    """Проверяет, находится ли путь в разрешенных директориях."""
    path = os.path.abspath(path)
    return any(path.startswith(os.path.abspath(allowed)) for allowed in allowed_paths)

def validate_file_path(file_path: str, should_exist: bool = True) -> tuple[bool, str]:
    """
    Проверяет валидность пути к файлу.
    
    Args:
        file_path: Путь к файлу
        should_exist: Должен ли файл существовать
        
    Returns:
        Кортеж (успех, сообщение об ошибке)
    """
    if not is_path_allowed(file_path):
        return False, f"Ошибка: Доступ к {file_path} запрещен. Разрешены только пути: {', '.join(allowed_paths)}"
    
    if should_exist and not os.path.exists(file_path):
        return False, f"Ошибка: Файл {file_path} не существует."
    
    return True, ""

@mcp.tool()
async def read_pdf(file_path: str, page_range: str = None, include_metadata: bool = True) -> str:
    """
    Читает содержимое PDF-файла и возвращает его текст.
    
    Args:
        file_path: Путь к PDF-файлу
        page_range: Диапазон страниц для извлечения (например, "1-5" или "2,4,6")
        include_metadata: Включать ли метаданные PDF в результат
    """
    try:
        valid, error_msg = validate_file_path(file_path)
        if not valid:
            return error_msg
        
        try:
            import PyPDF2
        except ImportError:
            return "Ошибка: Для работы с PDF требуется библиотека PyPDF2. Установите её с помощью команды: pip install PyPDF2"
        
        
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            
            info = reader.metadata
            total_pages = len(reader.pages)
            
            
            pages_to_extract = []
            if page_range:
                
                ranges = page_range.split(',')
                for r in ranges:
                    if '-' in r:
                        start, end = map(int, r.split('-'))
                        
                        start = max(1, min(start, total_pages))
                        end = max(1, min(end, total_pages))
                        
                        pages_to_extract.extend(range(start - 1, end))
                    else:
                        try:
                            page = int(r)
                            if 1 <= page <= total_pages:
                                pages_to_extract.append(page - 1)
                        except ValueError:
                            pass
                
                if not pages_to_extract:
                    pages_to_extract = range(total_pages)
            else:
                pages_to_extract = range(total_pages)
            
            result = []
            
            if include_metadata:
                result.append(f"=== ИНФОРМАЦИЯ О PDF ===")
                result.append(f"Название: {info.title if hasattr(info, 'title') and info.title else 'Не указано'}")
                result.append(f"Автор: {info.author if hasattr(info, 'author') and info.author else 'Не указан'}")
                result.append(f"Создан: {info.creation_date if hasattr(info, 'creation_date') and info.creation_date else 'Не указано'}")
                result.append(f"Изменен: {info.modification_date if hasattr(info, 'modification_date') and info.modification_date else 'Не указано'}")
                result.append(f"Количество страниц: {total_pages}")
                result.append("")
            
            result.append(f"=== СОДЕРЖИМОЕ ===")
            for i in sorted(pages_to_extract):
                if 0 <= i < total_pages:
                    page = reader.pages[i]
                    text = page.extract_text()
                    result.append(f"--- Страница {i + 1} ---")
                    if text:
                        result.append(text)
                    else:
                        result.append("[Страница не содержит текста или текст не может быть извлечен]")
                    result.append("")
            
            return "\n".join(result)
        
    except Exception as e:
        return f"Ошибка при чтении PDF-файла: {str(e)}"

@mcp.tool()
async def read_docx(file_path: str, format_type: str = "text", tables_only: bool = False) -> str:
    """
    Reads the contents of a DOCX file and returns its text.
    """
    try:
        valid, error_msg = validate_file_path(file_path)
        if not valid:
            return error_msg
        
        document = Document(file_path)
        
        if tables_only:
            return get_tables_info(document)
        
        if format_type.lower() == 'json':
            return get_document_as_json(document)
        
        return get_document_as_text(document)
    
    except Exception as e:
        return f"Ошибка при чтении DOCX-файла: {str(e)}"

def get_tables_info(document: Document) -> str:
    """Возвращает информацию о таблицах документа."""
    result = []
    result.append(f"Найдено таблиц: {len(document.tables)}")
    
    for t_idx, table in enumerate(document.tables):
        result.append(f"\nТаблица {t_idx+1}:")
        result.append(f"  Строк: {len(table.rows)}")
        result.append(f"  Столбцов: {len(table.columns) if table.rows else 0}")
        
        # Заголовок таблицы (если есть)
        if table.rows and len(table.rows) > 0:
            header_row = []
            for cell in table.rows[0].cells:
                header_row.append(cell.text.strip())
            result.append(f"  Заголовок: {' | '.join(header_row)}")
    
    return "\n".join(result)

def get_document_as_text(document: Document) -> str:
    """Возвращает содержимое документа в текстовом формате."""
    full_text = []
    
    paragraphs_text = []
    for i, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip():  
            paragraphs_text.append(f"[Абзац {i+1}] {paragraph.text}")
    
    if paragraphs_text:
        full_text.append("=== АБЗАЦЫ ===")
        full_text.extend(paragraphs_text)
    
    table_count = len(document.tables)
    if table_count > 0:
        full_text.append("\n=== ТАБЛИЦЫ ===")
        for t_idx, table in enumerate(document.tables):
            full_text.append(f"\n[Таблица {t_idx+1}]")
            for r_idx, row in enumerate(table.rows):
                row_text = []
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:  
                        row_text.append(f"({r_idx+1},{c_idx+1}): {cell_text}")
                if row_text:  
                    full_text.append(" | ".join(row_text))
    
    result = "\n".join(full_text)
    return result if result.strip() else "Документ пустой или не содержит текста."

def get_document_as_json(document: Document) -> str:
    """Возвращает содержимое документа в JSON формате."""
    result = {
        "paragraphs": [],
        "tables": []
    }
    
    # Абзацы
    for i, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip():
            result["paragraphs"].append({
                "index": i,
                "text": paragraph.text
            })
    
    # Таблицы
    for t_idx, table in enumerate(document.tables):
        table_data = {
            "index": t_idx,
            "rows": len(table.rows),
            "columns": len(table.columns) if table.rows else 0,
            "cells": []
        }
        
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    table_data["cells"].append({
                        "row": r_idx,
                        "column": c_idx,
                        "text": cell_text
                    })
        
        result["tables"].append(table_data)
    
    return json.dumps(result, ensure_ascii=False, indent=2)

@mcp.tool()
async def edit_docx(
    file_path: str, 
    replacements: Dict[str, str] = None,
    use_regex: bool = False,
    output_path: str = None,
    append_content: List[Dict[str, Any]] = None  
) -> str:
    """
    Edits a DOCX file by replacing the specified text fragments. 
    Note that to avoid creating a new file instead of the old one, 
    you need to explicitly say: "edit, but do not create a new one."

    Args:
        file_path: Путь к DOCX-файлу
        replacements: Словарь замен в формате {"старый_текст": "новый_текст"}
        use_regex: Использовать регулярные выражения (True/False)
        output_path: Путь для сохранения результата (если не указан, файл будет перезаписан)
        append_content: Список элементов для добавления в конец документа
            [
                {"type": "paragraph", "text": "Текст абзаца"},
                {"type": "heading", "text": "Заголовок", "level": 1},
                {"type": "table", "rows": [["ячейка1", "ячейка2"], ["ячейка3", "ячейка4"]]}
            ]
    """
    try:
        valid, error_msg = validate_file_path(file_path)
        if not valid:
            return error_msg
        
        
        if output_path:
            valid, error_msg = validate_file_path(output_path, should_exist=False)
            if not valid:
                return error_msg
        else:
            output_path = file_path
        
        document = Document(file_path)
        changes_count = 0
        
        #
        if replacements:
            
            if use_regex:
                
                compiled_patterns = []
                for pattern, replacement in replacements.items():
                    try:
                        compiled_patterns.append((re.compile(pattern), replacement))
                    except re.error as e:
                        return f"Ошибка в регулярном выражении '{pattern}': {str(e)}"
                
                
                changes_count = apply_regex_replacements(document, compiled_patterns)
            else:
                
                changes_count = apply_text_replacements(document, replacements)
        
        # Добавляем новый контент в конец документа, если указан
        if append_content:
            for item in append_content:
                item_type = item.get('type', 'paragraph')
                
                if item_type == 'paragraph':
                    text = item.get('text', '')
                    if text:
                        document.add_paragraph(text)
                        changes_count += 1
                
                elif item_type == 'heading':
                    text = item.get('text', '')
                    level = item.get('level', 1)
                    if text:
                        document.add_heading(text, level=level)
                        changes_count += 1
                
                elif item_type == 'table':
                    rows = item.get('rows', [])
                    if rows and len(rows) > 0 and len(rows[0]) > 0:
                        row_count = len(rows)
                        col_count = len(rows[0])
                        table = document.add_table(rows=row_count, cols=col_count)
                        
                        for i, row_data in enumerate(rows):
                            for j, cell_data in enumerate(row_data):
                                if j < col_count:  # Проверка на выход за границы
                                    cell = table.cell(i, j)
                                    cell.text = str(cell_data)
                        
                        changes_count += 1
                
                elif item_type == 'list':
                    items = item.get('items', [])
                    style = item.get('style', 'bullet')  # 'bullet' или 'number'
                    
                    if items:
                        for list_item in items:
                            if style == 'bullet':
                                document.add_paragraph(list_item, style='ListBullet')
                            else:
                                document.add_paragraph(list_item, style='ListNumber')
                        
                        changes_count += 1
        
        # Создаем директории для выходного пути, если необходимо
        output_dir = os.path.dirname(os.path.abspath(output_path))
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            
        # Сохраняем документ
        document.save(output_path)
        
        if changes_count > 0:
            return f"Файл {'сохранен как ' + output_path if output_path != file_path else file_path + ' обновлен'}. Выполнено изменений: {changes_count}."
        else:
            return f"В файле не было сделано изменений."
    
    except Exception as e:
        return f"Ошибка при редактировании DOCX-файла: {str(e)}"

def apply_text_replacements(document: Document, replacements: Dict[str, str]) -> int:
    """Применяет прямые замены текста в документе."""
    changes_count = 0
    
    # Заменяем текст в абзацах
    for paragraph in document.paragraphs:
        original_text = paragraph.text
        modified_text = original_text
        
        for old_text, new_text in replacements.items():
            if old_text in modified_text:
                modified_text = modified_text.replace(old_text, new_text)
        
        if original_text != modified_text:
            paragraph.clear()
            paragraph.add_run(modified_text)
            changes_count += 1
    
    # Заменяем текст в таблицах
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    modified_text = original_text
                    
                    for old_text, new_text in replacements.items():
                        if old_text in modified_text:
                            modified_text = modified_text.replace(old_text, new_text)
                    
                    if original_text != modified_text:
                        paragraph.clear()
                        paragraph.add_run(modified_text)
                        changes_count += 1
    
    return changes_count

def apply_regex_replacements(document: Document, compiled_patterns: List[tuple]) -> int:
    """Применяет замены с использованием регулярных выражений."""
    changes_count = 0
    
    # Заменяем текст в абзацах
    for paragraph in document.paragraphs:
        original_text = paragraph.text
        modified_text = original_text
        
        for pattern, replacement in compiled_patterns:
            modified_text = pattern.sub(replacement, modified_text)
        
        if original_text != modified_text:
            paragraph.clear()
            paragraph.add_run(modified_text)
            changes_count += 1
    
    # Заменяем текст в таблицах
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    modified_text = original_text
                    
                    for pattern, replacement in compiled_patterns:
                        modified_text = pattern.sub(replacement, modified_text)
                    
                    if original_text != modified_text:
                        paragraph.clear()
                        paragraph.add_run(modified_text)
                        changes_count += 1
    
    return changes_count

@mcp.tool()
async def edit_docx_table(file_path: str, table_index: int, operations: List[Dict[str, Any]], output_path: str = None, show_structure: bool = False, dry_run: bool = False) -> str:
    """
    Edits a DOCX file table.
    
    Args:
        file_path: Путь к DOCX-файлу
        table_index: Индекс таблицы (начиная с 0)
        operations: Список операций над таблицей
        output_path: Путь для сохранения результата (если не указан, файл будет перезаписан)
        show_structure: Показать структуру таблицы перед редактированием
        dry_run: Режим проверки без внесения изменений
    """
    try:
        valid, error_msg = validate_file_path(file_path)
        if not valid:
            return error_msg
            
        # Проверка выходного пути, если указан
        if output_path and not dry_run:
            valid, error_msg = validate_file_path(output_path, should_exist=False)
            if not valid:
                return error_msg
        else:
            output_path = file_path
            
        document = Document(file_path)
        
        if not document.tables or table_index >= len(document.tables):
            return f"Ошибка: Таблица с индексом {table_index} не найдена. Всего таблиц: {len(document.tables)}."
        
        table = document.tables[table_index]
        
        # Если запрошено, показываем структуру таблицы
        if show_structure:
            structure_info = []
            structure_info.append(f"Структура таблицы {table_index}:")
            structure_info.append(f"Количество строк: {len(table.rows)}")
            structure_info.append(f"Количество столбцов: {len(table.columns) if table.rows else 0}")
            
            # Показываем индексы и содержимое ячеек в виде таблицы
            headers = []
            for c_idx in range(len(table.columns)):
                headers.append(f"Col {c_idx}")
            structure_info.append("\n| Row # | " + " | ".join(headers) + " |")
            structure_info.append("|" + "-" * 7 + "|" + "".join(["-" * (len(h) + 2) + "|" for h in headers]))
            
            for r_idx, row in enumerate(table.rows):
                row_content = []
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip() if cell.text.strip() else "[пусто]"
                    # Ограничиваем длину для удобства чтения
                    if len(cell_text) > 15:
                        cell_text = cell_text[:12] + "..."
                    row_content.append(cell_text)
                structure_info.append(f"| {r_idx:5d} | " + " | ".join(row_content) + " |")
            
            return "\n".join(structure_info)
        
        if dry_run:
            results = []
            results.append("Предварительная проверка операций (без внесения изменений):")
            
            for op_idx, op in enumerate(operations):
                try:
                    op_type = op.get("type", "cell")
                    
                    if op_type == "cell":
                        row_idx = op.get("row", 0)
                        col_idx = op.get("column", 0)
                        new_text = op.get("text", "")
                        
                        if row_idx < 0 or row_idx >= len(table.rows):
                            results.append(f"❌ Операция {op_idx}: Некорректный индекс строки {row_idx}. Допустимые значения: 0-{len(table.rows)-1}")
                            continue
                        
                        if col_idx < 0 or col_idx >= len(table.rows[0].cells):
                            results.append(f"❌ Операция {op_idx}: Некорректный индекс столбца {col_idx}. Допустимые значения: 0-{len(table.rows[0].cells)-1}")
                            continue
                        
                        cell = table.cell(row_idx, col_idx)
                        current_text = cell.text.strip()
                        results.append(f"✓ Операция {op_idx}: Изменение ячейки ({row_idx},{col_idx})")
                        results.append(f"   Текущий текст: '{current_text}'")
                        results.append(f"   Новый текст: '{new_text}'")
                    
                    elif op_type == "row_range":
                        row_start = op.get("row_start", 0)
                        row_end = op.get("row_end", 0)
                        col_idx = op.get("column", 0)
                        new_text = op.get("text", "")
                        
                        if row_start < 0 or row_start >= len(table.rows) or row_end < 0 or row_end >= len(table.rows):
                            results.append(f"❌ Операция {op_idx}: Некорректный диапазон строк {row_start}-{row_end}. Допустимые значения: 0-{len(table.rows)-1}")
                            continue
                        
                        if col_idx < 0 or col_idx >= len(table.rows[0].cells):
                            results.append(f"❌ Операция {op_idx}: Некорректный индекс столбца {col_idx}. Допустимые значения: 0-{len(table.rows[0].cells)-1}")
                            continue
                        
                        results.append(f"✓ Операция {op_idx}: Изменение ячеек в строках {row_start}-{row_end}, столбец {col_idx}")
                        for row_idx in range(row_start, row_end + 1):
                            cell = table.cell(row_idx, col_idx)
                            current_text = cell.text.strip()
                            results.append(f"   Строка {row_idx}: '{current_text}' -> '{new_text}'")
                    
                    elif op_type == "column":
                        col_idx = op.get("column", 0)
                        new_text = op.get("text", "")
                        
                        if col_idx < 0 or col_idx >= len(table.columns):
                            results.append(f"❌ Операция {op_idx}: Некорректный индекс столбца {col_idx}. Допустимые значения: 0-{len(table.columns)-1}")
                            continue
                        
                        results.append(f"✓ Операция {op_idx}: Изменение всех ячеек в столбце {col_idx}")
                        for row_idx in range(len(table.rows)):
                            cell = table.cell(row_idx, col_idx)
                            current_text = cell.text.strip()
                            results.append(f"   Строка {row_idx}: '{current_text}' -> '{new_text}'")
                    
                    else:
                        results.append(f"❌ Операция {op_idx}: Неизвестный тип операции '{op_type}'")
                
                except Exception as e:
                    results.append(f"❌ Операция {op_idx}: Ошибка при проверке: {str(e)}")
            
            return "\n".join(results)
        
        # Реальное редактирование
        changes_made = []
        errors = []
        changes_count = 0
        
        for op_idx, op in enumerate(operations):
            try:
                op_type = op.get("type", "cell")
                
                if op_type == "cell":
                    # Изменение одной ячейки
                    row_idx = op.get("row", 0)
                    col_idx = op.get("column", 0)
                    new_text = op.get("text", "")
                    
                    if row_idx < 0 or row_idx >= len(table.rows):
                        errors.append(f"Операция {op_idx}: Некорректный индекс строки {row_idx}. Допустимые значения: 0-{len(table.rows)-1}")
                        continue
                    
                    if col_idx < 0 or col_idx >= len(table.rows[0].cells):
                        errors.append(f"Операция {op_idx}: Некорректный индекс столбца {col_idx}. Допустимые значения: 0-{len(table.rows[0].cells)-1}")
                        continue
                    
                    cell = table.cell(row_idx, col_idx)
                    if update_cell_text(cell, new_text):
                        changes_count += 1
                        changes_made.append(f"Изменена ячейка ({row_idx},{col_idx}): '{cell.text.strip()}' -> '{new_text}'")
                
                elif op_type == "row_range":
                    # Изменение диапазона ячеек
                    row_start = op.get("row_start", 0)
                    row_end = op.get("row_end", 0)
                    col_idx = op.get("column", 0)
                    new_text = op.get("text", "")
                    
                    if row_start < 0 or row_start >= len(table.rows) or row_end < 0 or row_end >= len(table.rows):
                        errors.append(f"Операция {op_idx}: Некорректный диапазон строк {row_start}-{row_end}. Допустимые значения: 0-{len(table.rows)-1}")
                        continue
                    
                    if col_idx < 0 or col_idx >= len(table.rows[0].cells):
                        errors.append(f"Операция {op_idx}: Некорректный индекс столбца {col_idx}. Допустимые значения: 0-{len(table.rows[0].cells)-1}")
                        continue
                    
                    for row_idx in range(row_start, row_end + 1):
                        cell = table.cell(row_idx, col_idx)
                        if update_cell_text(cell, new_text):
                            changes_count += 1
                            changes_made.append(f"Изменена ячейка ({row_idx},{col_idx}): '{cell.text.strip()}' -> '{new_text}'")
                
                elif op_type == "column":
                    # Изменение всего столбца
                    col_idx = op.get("column", 0)
                    new_text = op.get("text", "")
                    skip_header = op.get("skip_header", False)  # Опция пропуска заголовка
                    
                    if col_idx < 0 or col_idx >= len(table.columns):
                        errors.append(f"Операция {op_idx}: Некорректный индекс столбца {col_idx}. Допустимые значения: 0-{len(table.columns)-1}")
                        continue
                    
                    start_row = 1 if skip_header else 0  # Пропускаем первую строку, если это заголовок
                    
                    for row_idx in range(start_row, len(table.rows)):
                        cell = table.cell(row_idx, col_idx)
                        if update_cell_text(cell, new_text):
                            changes_count += 1
                            changes_made.append(f"Изменена ячейка ({row_idx},{col_idx}): '{cell.text.strip()}' -> '{new_text}'")
            
            except Exception as e:
                errors.append(f"Ошибка в операции {op_idx}: {str(e)}")
        
        # Отчет о результатах
        result_message = []
        
        if changes_count > 0:
            # Создаем директории для выходного пути, если необходимо
            output_dir = os.path.dirname(os.path.abspath(output_path))
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
                
            # Сохраняем документ
            document.save(output_path)
            
            result_message.append(f"Таблица {table_index} в файле {'сохранена как ' + output_path if output_path != file_path else file_path + ' обновлена'}.")
            result_message.append(f"Изменено ячеек: {changes_count}")
            if changes_made:
                result_message.append("\nВнесенные изменения:")
                result_message.extend(changes_made)
        else:
            result_message.append(f"В таблице {table_index} не было сделано изменений.")
        
        if errors:
            result_message.append("\nОшибки при выполнении:")
            result_message.extend(errors)
        
        return "\n".join(result_message)
    
    except Exception as e:
        return f"Ошибка при редактировании таблицы в DOCX-файле: {str(e)}"

def update_cell_text(cell, new_text):
    """
    Безопасно обновляет текст в ячейке таблицы, сохраняя структуру.
    
    Args:
        cell: Ячейка таблицы
        new_text: Новый текст для ячейки
    """
    # Проверка на изменение содержимого
    if cell.text.strip() == new_text.strip():
        return False  # Изменений нет
    
    # Более безопасный способ обновления текста без очистки всех параграфов
    if len(cell.paragraphs) > 0:
        # Обновляем первый параграф
        first_paragraph = cell.paragraphs[0]
        
        # Сохраняем стиль и форматирование
        original_style = first_paragraph.style
        
        # Очищаем только текст в первом параграфе, сохраняя его структуру
        for run in list(first_paragraph.runs):
            run.text = ""
        
        # Добавляем новый текст
        first_paragraph.add_run(new_text)
        
        # Восстанавливаем стиль
        first_paragraph.style = original_style
        
        # Удаляем лишние параграфы, если они есть
        if len(cell.paragraphs) > 1:
            for p in cell.paragraphs[1:]:
                p.text = ""
    else:
        # Если нет параграфов, создаем новый
        cell.add_paragraph(new_text)
    
    return True  # Изменения сделаны

@mcp.tool()
async def create_docx(file_path: str, content: List[Dict[str, Any]], template_path: str = None) -> str:
    """
    Creates a new DOCX file with the specified content. 
    Available elements include: paragraphs, headings, and tables. 
    Note that editing is done in markdown format, 
    and some elements may appear differently than intended.

    """
    try:
        valid, error_msg = validate_file_path(file_path, should_exist=False)
        if not valid:
            return error_msg
        
        # Если указан шаблон, проверяем его существование
        if template_path:
            valid, error_msg = validate_file_path(template_path)
            if not valid:
                return error_msg
            document = Document(template_path)
        else:
            document = Document()
        
        for item in content:
            item_type = item.get('type', 'paragraph')
            
            if item_type == 'paragraph':
                text = item.get('text', '')
                if text:
                    document.add_paragraph(text)
            
            elif item_type == 'heading':
                text = item.get('text', '')
                level = item.get('level', 1)
                if text:
                    document.add_heading(text, level=level)
            
            elif item_type == 'table':
                rows = item.get('rows', [])
                if rows and len(rows) > 0 and len(rows[0]) > 0:
                    row_count = len(rows)
                    col_count = len(rows[0])
                    table = document.add_table(rows=row_count, cols=col_count)
                    
                    for i, row_data in enumerate(rows):
                        for j, cell_data in enumerate(row_data):
                            if j < col_count:  # Проверка на выход за границы
                                cell = table.cell(i, j)
                                cell.text = str(cell_data)
            
            elif item_type == 'list':
                items = item.get('items', [])
                style = item.get('style', 'bullet')  # 'bullet' или 'number'
                
                for list_item in items:
                    if style == 'bullet':
                        document.add_paragraph(list_item, style='ListBullet')
                    else:
                        document.add_paragraph(list_item, style='ListNumber')
        
        # Создаем директории, если их нет
        directory = os.path.dirname(os.path.abspath(file_path))
        if directory and not os.path.exists(directory):
            os.makedirs(directory, exist_ok=True)
        
        # Сохраняем документ
        document.save(file_path)
        
        # Проверяем, что файл был создан
        if not os.path.exists(file_path):
            return f"Ошибка: Не удалось создать файл {file_path}."
        
        # Проверяем, что файл не пустой
        if os.path.getsize(file_path) == 0:
            return f"Предупреждение: Файл {file_path} создан, но имеет нулевой размер."
        
        return f"Файл {file_path} успешно создан."
    
    except Exception as e:
        return f"Ошибка при создании DOCX-файла: {str(e)}"

# Запускаем сервер
if __name__ == "__main__":
    print(f"DOCX сервер запущен. Разрешенные пути: {allowed_paths}")
    mcp.run(transport='stdio')
